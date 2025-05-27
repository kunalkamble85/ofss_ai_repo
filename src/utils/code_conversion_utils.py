import traceback
import pandas as pd
import os
import tiktoken
from oci_utils import generate_oci_gen_ai_response, generate_embeddings
import re
import json
import database_utils
import logging
import prompts
from io import BytesIO
from docx import Document
from markdown import markdown
from bs4 import BeautifulSoup
import base64
import time


conversion_folder = "code_conversion_files"
react_project_zip_file = f"{conversion_folder}/target_react_structure/react_project.zip"
FILE_EXCLUSIONS = ["png","jpg","jpeg","txt","md","json","xml","csv","zip","gz","tar","pdf","doc","docx","xls","xlsx","ppt","pptx","pyc","exe","dll","so","lib",
                   "obj","bin","jar","war","ear","class","vbproj","csproj","vcxproj","vcproj","filters","user","data","vspscc","sln","suo","webinfo","xsd","xslt"]
ANGULER_JS_FILE_INCLUSIONS = ["png","jpg","jpeg","ts","js","html","css","xml","csv"]
ROUTER_FILE_FIND_STRING = "$urlRouterProvider"
SAMPLE_REACT_PROJECT_PATH = "react_project"
TEST_SIZE = 3
INCLUDE_ALL_DOCUMENT_BY_DEFAULT = True
REVIEW_LOOP_COUNT = 1

def __get_details_from_repsonse(response):
    logging.warning("START: __get_details_from_repsonse")
    """
    Extracts details such as target file name, folder path, and converted code from the response.

    Args:
        response (str): The response string containing the details.

    Returns:
        tuple: A tuple containing target file name, folder path, and converted code.
    """
    target_file_name = re.search(r"<target_file_name>(.*?)</target_file_name>", response, re.DOTALL)
    if target_file_name:
        target_file_name =  target_file_name.group(1)
    logging.warning(f"target_file_name:{target_file_name}") # AccountCtrl.jsx

    target_folder_path = re.search(r"<target_folder_path>(.*?)</target_folder_path>", response, re.DOTALL)
    if target_folder_path:
        target_folder_path =  target_folder_path.group(1)
    logging.warning(f"target_folder_path:{target_folder_path}") # AccountCtrl.jsx

    converted_code = re.search(r"<converted_code>(.*?)</converted_code>", response, re.DOTALL)
    if converted_code:
        converted_code =  converted_code.group(1)
    # logging.warning(converted_code) 
    logging.warning("END: __get_details_from_repsonse")
    return target_file_name, target_folder_path, converted_code


def get_view_controller_info(files_to_process, routing_information):
    logging.warning("START: get_view_controller_info")
    """
    Maps view files to their corresponding controller files based on routing information.

    Args:
        files_to_process (dict): Dictionary of files to process.
        routing_information (list): List of routing information.

    Returns:
        dict: A mapping of view files to controller files.
    """
    logging.warning(f"routing_information:{routing_information}")
    view_controller_mapping = {}
    for route_info in routing_information:
        view_path =  route_info["View_Path"]
        controller = route_info["Controller"]
        controller = controller.replace("Ctrl", "")
        gotController = False
        gotView = False
        for _, (file_name, _) in files_to_process.items():
            fname = file_name.split("/")[-1]
            if "controller" in file_name and controller in fname:
                controller = file_name
                gotController = True
                break
        for _, (file_name, _) in files_to_process.items():
            if file_name.endswith(view_path):
                view_path = file_name
                gotView = True
                break
        if gotController and gotView:
            view_controller_mapping[view_path] = controller
    logging.warning(f"view_controller_mapping:{view_controller_mapping}")
    logging.warning("END: get_view_controller_info")
    return view_controller_mapping


def generate_route_file(llm_model, files_to_process):
    logging.warning("START: generate_route_file")
    """
    Generates the ReactJS route file from AngularJS routing information.

    Args:
        llm_model (str): The LLM model to use for generating the route file.
        files_to_process (dict): Dictionary of files to process.

    Returns:
        tuple: Details about the generated route file and view-controller mapping.
    """
    user_request_details_id, file_name, file_content, routing_information_prompt = prompts.get_routing_information_prompt(files_to_process)

    view_controller_info ={}
    if user_request_details_id is not None:
        response = generate_oci_gen_ai_response(llm_model, [{"role":"user", "content": routing_information_prompt}])  
        response = response.replace("```json","").replace("```","")
        routing_information = json.loads(response)
        view_controller_info = get_view_controller_info(files_to_process, routing_information)    
        code_converion_prompt = prompts.get_ang_to_react_code_conversion_prompt(files_to_process, file_name,  file_content, None)
        response = generate_oci_gen_ai_response(llm_model, [{"role":"user", "content": code_converion_prompt}])                                     
        _, _, converted_code = __get_details_from_repsonse(response)
        target_file_to_create = "react_project/src/App.jsx"
        logging.warning(f"target_file_to_create:{target_file_to_create}")
        converted_code = converted_code.replace("```javascript","").replace("```","").replace("```jsx","").replace("```","").replace("jsx","")
        logging.warning("END: generate_route_file")
        return user_request_details_id, file_name, target_file_to_create, view_controller_info, converted_code
    else:
        logging.warning("END: generate_route_file")
        return None, None, None, None, None


def handle_view_controller_files(llm_model, files_to_process, file_objects, target_folder_structure, view_controller_info, route_file_content, test_mode):
    logging.warning("START: handle_view_controller_files")
    """
    Handles the conversion of view and controller files from AngularJS to ReactJS.

    Args:
        llm_model (str): The LLM model to use for conversion.
        files_to_process (dict): Dictionary of files to process.
        file_objects (list): List of file objects from the database.
        target_folder_structure (list): List of target folder structure.
        view_controller_info (dict): Mapping of view files to controller files.
        route_file_content (str): Content of the route file.
        test_mode (bool): Whether the function is running in test mode.

    Returns:
        tuple: Output, errors, and updated target folder structure.
    """
    output = []
    errors = []
    for view_file_name, controller_file_name in view_controller_info.items():
        logging.warning(f"Processing view file: {view_file_name}, controller_file:{controller_file_name}")            
        try:
            view_controller_conversion_prompt = prompts.get_view_controller_conversion_prompt(files_to_process, target_folder_structure, view_file_name, controller_file_name, route_file_content, file_objects)
            logging.warning(f"code_converion_prompt:{view_controller_conversion_prompt}")
            response = generate_oci_gen_ai_response(llm_model, [{"role":"user", "content": view_controller_conversion_prompt}])                                     
            target_file_name, target_folder_path, converted_code = __get_details_from_repsonse(response)
            target_file_to_create = SAMPLE_REACT_PROJECT_PATH + "/" + target_folder_path + "/" + target_file_name            
            target_file_to_create = target_file_to_create.replace("\\","/").replace("//","/").replace("/\\","/").replace("react_project/react_project", "react_project")            
            logging.warning(f"target_file_to_create:{target_file_to_create}")
            converted_code = converted_code.replace("```javascript","").replace("```","").replace("```jsx","").replace("```","").replace("jsx","")            
            output.append([view_file_name, target_file_to_create, converted_code, "Y", None])
            output.append([controller_file_name, target_file_to_create, converted_code, "Y", None])
            target_folder_structure.append(target_file_to_create)                
        except Exception as e:      
            error = str(e)          
            logging.warning(traceback.format_exc())
            errors.append(view_file_name)
            errors.append(controller_file_name)
            output.append([view_file_name, None, None, "N", error])
            output.append([controller_file_name, None, None, "N", error])
        if test_mode:
            break
    logging.warning("END: handle_view_controller_files")
    return output, errors, target_folder_structure


def get_number_of_tokens(file_content):
    logging.warning("START: get_number_of_tokens")
    """
    Calculates the number of tokens in the given file content.

    Args:
        file_content (str): The content of the file.

    Returns:
        int: The number of tokens in the file content.
    """
    encoding = tiktoken.get_encoding("cl100k_base")
    tokens = encoding.encode(file_content)
    logging.warning("END: get_number_of_tokens")
    return len(tokens)

def get_files_from_database(user_request_id, exclude_files=FILE_EXCLUSIONS, include_files=None, include_empty_folders = False, skip_test_files = True):
    logging.warning("START: get_files_from_database")
    """
    Fetches files from the database for a given user request ID.

    Args:
        user_request_id (int): The ID of the user request.
        exclude_files (list): List of file extensions to exclude.
        include_files (list): List of file extensions to include.
        include_empty_folders (bool): Whether to include empty folders.

    Returns:
        tuple: Files to process, directories, and file objects.
    """
    db_get_user_request_details = database_utils.db_get_user_request_details(user_request_id)
    files_to_process = {}
    dirs = []
    file_objects = db_get_user_request_details["items"]    
    for item in file_objects:
        user_request_details_id = item["user_request_details_id"]
        file_name = item["file_name"]
        if skip_test_files and "/tests/" in file_name: 
            logging.warning(f"Skipping test file: {file_name}")
            continue
        extension = file_name.split(".")[-1]
        if "." not in file_name or (exclude_files is None or extension not in exclude_files) and (include_files is None or extension in include_files):            
            file_content = item["file_content"]
            files_to_process[user_request_details_id] = (file_name, file_content)
    if include_empty_folders and file_name.endswith("/"):
        dirs.append(file_name)
    logging.warning("END: get_files_from_database")
    return files_to_process, dirs, file_objects

def get_files_from_directory(folder_path, exclude_files = None, include_files = None, include_empty_folders = False):
    logging.warning("START: get_files_from_directory")
    """
    Fetches files from a directory based on inclusion and exclusion criteria.

    Args:
        folder_path (str): The path of the folder.
        exclude_files (list): List of file extensions to exclude.
        include_files (list): List of file extensions to include.
        include_empty_folders (bool): Whether to include empty folders.

    Returns:
        list: List of filtered files.
    """
    filtered_files = []
    files = [val for sublist in [[os.path.join(i[0], j) for j in i[2]] for i in os.walk(folder_path)] for val in sublist]
    for file in files:
        extension = file.split(".")[-1]
        if "." not in file or (exclude_files is None or extension not in exclude_files) and (include_files is None or extension in include_files):
            file_name = file.replace("\\","/")
            filtered_files.append(file_name)
    if include_empty_folders:
        dirs = [i[0] for i in os.walk(folder_path)]
        filtered_files = filtered_files + dirs
    logging.warning("END: get_files_from_directory")
    return filtered_files

    
def generate_analysis_report(user_request_id, test_mode):
    """
    Generates an analysis report for the files associated with a given user request ID.

    Args:
        user_request_id (int): The ID of the user request.
        test_mode (bool): Whether the function is running in test mode.

    Raises:
        Exception: If there is an error during the analysis process.
    """
    logging.warning(f"START: generate_analysis_report with user_request_id={user_request_id}")
    start_time = time.time()
    try:
        db_user_request = database_utils.db_get_user_request(user_request_id)
        if db_user_request["analysis_request_status"] == "In-Progress":
            logging.warning(f"Analysis process is already running for the request ID:{user_request_id}")
            return
                
        database_utils.update_user_request_status(user_request_id, {"analysis_request_status":"In-Progress","error_message":None})
        package_name = db_user_request["zip_file_name"].split(".")[0]
        updated_by = db_user_request["updated_by"]
        logging.warning(f"Package name: {package_name}, Updated by: {updated_by}")
        
        files_to_process, _, _= get_files_from_database(user_request_id, exclude_files=FILE_EXCLUSIONS, include_files=None)
        logging.warning(f"Number of files to process: {len(files_to_process)}")
        
        if len(files_to_process) == 0:
            logging.warning("No files found in the source folder")
            raise Exception("No files found in the source folder")
        
        errors = []
        df = pd.DataFrame(columns=['user_request_details_id', 'user_request_id', 'num_lines', 'num_lines_no_docs', 'token_count', 'token_count_no_docs', 'success_flag', 'error_details', 'updated_by'])
        data_counter = 0
        
        for user_request_details_id, (file_name, file_content) in files_to_process.items():
            try:
                logging.warning(f"Processing file: {file_name}")
                total_tokens = get_number_of_tokens(file_content)
                tokens_without_comments = number_of_lines = 0
                number_of_lines_without_comments = 0
                tokens_without_comments_content = ""
                
                for line in file_content.splitlines():
                    number_of_lines += 1
                    if not line.startswith("#") and line.strip() != "" and not line.startswith("//"):
                        number_of_lines_without_comments += 1
                        tokens_without_comments_content += line
                
                tokens_without_comments = get_number_of_tokens(tokens_without_comments_content)
                df.loc[data_counter] = [user_request_details_id, user_request_id, number_of_lines, number_of_lines_without_comments, total_tokens, tokens_without_comments, "Y", None, updated_by]
                data_counter += 1
                logging.warning(f"File processed successfully: {file_name}")
                if test_mode and data_counter >= TEST_SIZE:
                    break
            except Exception as e:
                logging.warning(f"Error processing file {file_name}: {str(e)}")
                logging.warning(traceback.format_exc())
                errors.append(file_name)
        database_utils.update_user_request_status(user_request_id, {"analysis_request_status":"Done","error_message":None})
        logging.warning(f"Errors encountered: {errors}")
        logging.warning(f"Completed generating analysis report for package_name: {package_name}")
        database_utils.update_analysis_report(df)
    except Exception as e:
        logging.warning(f"Error in generate_analysis_report: {str(e)}")
        logging.warning(traceback.format_exc())
        database_utils.update_user_request_status(user_request_id, {"analysis_request_status":"Done", "error_message":str(e)})
        raise
    finally:
        end_time = time.time()
        total_time = end_time - start_time
        logging.warning(f"END: generate_analysis_report with user_request_id={user_request_id}. Total time taken: {total_time:.2f} seconds")


def split_documentations(documentation_content_list):
    logging.warning("START: split_documentations")
    """
    Splits a list of documentation strings into multiple strings, each within the 100,000 token limit.

    Args:
        documentation_content_list (list): List of documentation strings.

    Returns:
        list: List of strings, each within the token limit.
    """
    max_tokens = 100000
    result = []
    current_chunk = []
    current_tokens = 0
    for doc in documentation_content_list:
        doc_tokens = get_number_of_tokens(doc)
        if current_tokens + doc_tokens > max_tokens:
            if current_chunk:
                result.append("\n".join(current_chunk))
            current_chunk = [doc]
            current_tokens = doc_tokens
        else:
            current_chunk.append(doc)
            current_tokens += doc_tokens
    if current_chunk:
        result.append("\n".join(current_chunk))
    logging.warning("END: split_documentations")
    return result


def run_doc_generation_in_feedback_loop(component_doc_name, generation_prompt, llm_model, documentation_content, additional_context):
    logging.warning("START: run_doc_generation_in_feedback_loop")
    doc_chat_messages = []
    doc_chat_messages.append({"role": "user", "content": generation_prompt})
    review_count = 0
    component_content = ""
    while True:                
        component_content = generate_oci_gen_ai_response(llm_model, doc_chat_messages)
        doc_chat_messages.append({"role": "assistant", "content": component_content})
        # Do review         
        if review_count == REVIEW_LOOP_COUNT: break        
        else: review_count = review_count + 1 
        logging.warning(f"Running Documentation Review for: {component_doc_name}")
        review_prompt = prompts.get_review_prompt(documentation_content, additional_context, component_content)
        review_comments = generate_oci_gen_ai_response(llm_model, [{"role": "user", "content": review_prompt}])
        generation_prompt_with_review = f"""Below is my review comments, please finetune your response accordingly. 
        Do not put any unnecessary commentary. Provide answer in originally asked format. 
        ### Review Comments
        {review_comments}
        """
        doc_chat_messages.append({"role": "user", "content": generation_prompt_with_review})
    logging.warning("END: run_doc_generation_in_feedback_loop")
    return component_content


def generate_brd_report(llm_model, documentation_content_list, user_documentation_components, additional_context, existing_analysis_report_list):
    logging.warning("START: generate_brd_report")    
    db_documentation_components_details = database_utils.db_get_documentation_components()
    """
    Generates a Business Requirement Document (BRD) based on the provided documentation content.
    """
    logging.warning(f"Generating BRD content for the documentation.")

    documentation_components = []

    if len(user_documentation_components) == 0 and INCLUDE_ALL_DOCUMENT_BY_DEFAULT:
            for db_documentation_components_detail in db_documentation_components_details:
                documentation_components.append((db_documentation_components_detail["target_file_name"], 
                                            db_documentation_components_detail["documentation_component"], db_documentation_components_detail["prompt"]))
    else:
        for user_documentation_component_id in user_documentation_components:
            for db_documentation_components_detail in db_documentation_components_details:
                if user_documentation_component_id == db_documentation_components_detail["documentation_components_id"]:
                    documentation_components.append((db_documentation_components_detail["target_file_name"], 
                                        db_documentation_components_detail["documentation_component"], db_documentation_components_detail["prompt"]))
                    break
    
    brds = []
    brd_contents = ""
    splitted_documentation_content_list = split_documentations(documentation_content_list)
    for component_doc_name, component, component_prompt in documentation_components:
        logging.warning(f"Generating BRD content for {component}.")
        documentation_already_generated = False
        for _, _, documentation_file_name, documentation_file_content in existing_analysis_report_list:
            if component_doc_name == documentation_file_name:
                documentation_already_generated = True
                break
        if documentation_already_generated:
            logging.warning(f"Documentation is already generated for {component}.")
            # Commented otherwise it will save again into database
            # brds.append((component_doc_name, documentation_file_content))
            brd_contents = f"{brd_contents}\n\n{documentation_file_content}" 
        else:
            component_contents = [] 
            counter = 1       
            for documentation_content in splitted_documentation_content_list:
                logging.warning(f"Started documentating split for counter: {counter}")
                counter = counter + 1
                if "BusinessRequirementDocument" in component_doc_name:
                    generation_prompt = prompts.get_brd_generation_prompt(documentation_content, additional_context)
                else:
                    generation_prompt = component_prompt.replace("$DOCUMENTATION$", documentation_content)
                
                component_content = run_doc_generation_in_feedback_loop(component_doc_name, generation_prompt, llm_model,documentation_content, additional_context)
                component_contents.append(component_content)  
            if len(component_contents) > 1:
                component_content_reponse = "\n".join(component_contents)
                document_consolidation_prompt = prompts.get_document_consolidation_prompt(component_content_reponse, additional_context)                  
                component_content = generate_oci_gen_ai_response(llm_model, [{"role": "user", "content": document_consolidation_prompt}])
                brd_contents = f"{brd_contents}\n\n{component_content}" 
            else:
                brd_contents = f"{brd_contents}\n\n{component_contents[0]}"
            brds.append((component_doc_name, "\n\n\n".join(component_contents)))

    logging.warning("END: generate_brd_report")
    return brds, brd_contents


def md_to_docx_content(md_content):
    logging.warning("START: md_to_docx_content")
    """
    Converts a markdown string to a docx Document object and returns the binary content.

    Args:
        md_content (str): Markdown formatted string.

    Returns:
        bytes: Binary content of the generated docx file.
    """
    # Convert markdown to HTML
    html = markdown(md_content, extensions=['fenced_code', 'tables'])
    soup = BeautifulSoup(html, "html.parser")

    doc = Document()

    for element in soup.children:
        if element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Number')
        elif element.name == 'pre':
            code = element.get_text()
            doc.add_paragraph(code, style='Intense Quote')
        elif element.name == 'p':
            doc.add_paragraph(element.get_text())
        elif element.name == 'table':
            rows = element.find_all('tr')
            if rows:
                cols = rows[0].find_all(['td', 'th'])
                table = doc.add_table(rows=len(rows), cols=len(cols))
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.get_text()

    # Save to BytesIO and return bytes
    output = BytesIO()
    doc.save(output)
    logging.warning("END: md_to_docx_content")
    return base64.b64encode(output.getvalue()).decode("utf-8")
    # with open("output.docx", "wb") as f:
    #     f.write(output.getvalue())
    # return output.getvalue()

def generate_docx_from_md(df):
    logging.warning("START: generate_docx_from_md")
    df['documentation_file_name_docx'] = df['documentation_file_name'].apply(lambda x: x.replace('.md', '.docx') if x.endswith('.md') else x + '.docx')
    df['documentation_file_content_docx'] = df['documentation_file_content'].apply(md_to_docx_content)
    logging.warning("END: generate_docx_from_md")
    return df


def generate_documentation_report(llm_model, user_request_id, test_mode, additional_context, documentation_components):
    logging.warning("START: generate_documentation_report")
    """
    Generates a documentation report for the files associated with a given user request ID.

    Args:
        llm_model (str): The LLM model to use for generating documentation.
        user_request_id (int): The ID of the user request.
        test_mode (bool): Whether the function is running in test mode.
        additional_context (str): Additional context to include in the documentation.

    Raises:
        Exception: If there is an error during the documentation generation process.
    """
    logging.warning(f"START: generate_documentation_report with user_request_id={user_request_id}, test_mode={test_mode}")
    start_time = time.time()
    try:
        db_user_request = database_utils.db_get_user_request(user_request_id)
        if db_user_request["documentation_request_status"] == "In-Progress":
            logging.warning(f"Documentation process is already running for the request ID:{user_request_id}")
            return
        database_utils.update_user_request_status(user_request_id, {"documentation_request_status":"In-Progress", "error_message":None})        
        existing_analysis_report_list = database_utils.db_get_documentation_and_file_details(user_request_id)

        package_name = db_user_request["zip_file_name"].split(".")[0]
        updated_by = db_user_request["updated_by"]
        source_language = database_utils.get_language_details(db_user_request["source_lang_id"])

        logging.warning(f"Package name: {package_name}, Updated by: {updated_by}, Source language: {source_language}")
        
        files_to_process, _, _ = get_files_from_database(user_request_id, exclude_files=FILE_EXCLUSIONS, include_files=None)
        logging.warning(f"Number of files to process: {len(files_to_process)}")

        if len(files_to_process) == 0:
            logging.warning("No files found in the source folder")
            raise Exception("No files found in the source folder")

        errors = []
        df = pd.DataFrame(columns=['user_request_details_id', 'user_request_id', 'documentation_file_name', 'documentation_file_content', 'success_flag', 'error_details', 'updated_by'])
        data_counter = 0
        documentation_content_list = []

        for user_request_details_id, (file_name, file_content) in files_to_process.items():
            try:
                logging.warning(f"Processing file: {file_name}")
                analysis_already_done = False
                for analysis_report_file_name, _, documentation_file_name, documentation_file_content in existing_analysis_report_list:
                    if analysis_report_file_name == file_name:
                        analysis_already_done = True
                        break
                if analysis_already_done:
                    logging.warning(f"Documentation is already generated for this file: {file_name}")
                    documentation_content_list.append(f"File:{documentation_file_name}\nContent of file:\n{documentation_file_content}")
                else:
                    document_generation_prompt = prompts.get_document_generation_prompt(source_language, files_to_process, file_name, file_content, additional_context)
                    response = generate_oci_gen_ai_response(llm_model, [{"role": "user", "content": document_generation_prompt}])
                                        
                    if "." in file_name:
                        extension = file_name.split(".")[-1]
                        target_file = file_name.replace(f".{extension}", "_documentation.md")
                    else:
                        target_file = file_name + "_documentation.md"

                    logging.warning(f"Generated documentation for file: {file_name}, Target file: {target_file}")
                    documentation_content_list.append(f"File:{target_file}\nContent of file:\n{response}")
                    df.loc[data_counter] = [user_request_details_id, user_request_id, target_file, response, "Y", None, updated_by]
                    data_counter += 1

                    if test_mode and data_counter >= TEST_SIZE:
                        logging.warning("Test mode enabled, stopping after processing 3 files.")
                        break
            except Exception as e:
                logging.warning(f"Error processing file {file_name}: {str(e)}")
                logging.warning(traceback.format_exc())
                errors.append(file_name)

        brd_reponse, brd_contents = generate_brd_report(llm_model, documentation_content_list, documentation_components, additional_context, existing_analysis_report_list)
        for (component_doc_name, component_content) in brd_reponse:
            if "BusinessRequirementDocument" in  component_doc_name: component_content = brd_contents
            df.loc[data_counter] = [None, user_request_id, component_doc_name, component_content, "Y", None, updated_by]
            data_counter += 1
        logging.warning(f"Errors encountered: {errors}")
        logging.warning(f"Completed generating documentation report for package_name: {package_name}")
        df = generate_docx_from_md(df)
        database_utils.update_documentation_report(df)
        database_utils.update_user_request_status(user_request_id, {"documentation_request_status":"Done","error_message":None})
    except Exception as e:
        logging.warning(f"Error in generate_documentation_report: {str(e)}")
        logging.warning(traceback.format_exc())
        database_utils.update_user_request_status(user_request_id, {"documentation_request_status":"Done", "error_message":str(e)})
        raise
    finally:
        end_time = time.time()
        total_time = end_time - start_time
        logging.warning(f"END: generate_documentation_report with user_request_id={user_request_id}, test_mode={test_mode}. Total time taken: {total_time:.2f} seconds")


def handle_documentation_user_request(llm_model, user_request_id, additional_context, user_query, use_embedding):
    logging.warning("START: handle_documentation_user_request")
    try:        
        code_documentation_details = database_utils.db_get_documentation_and_file_details(user_request_id)
        if len(code_documentation_details) == 0:
            return "Documentation is not generated for the code. Please generate the documentation first and then re-run the user query."
    
        if use_embedding:
            embedding_pending_docs = database_utils.db_get_documentation_for_embedding()
            embedding_pending_docs = [entry for entry in embedding_pending_docs if entry.get("documentation_file_content") and entry.get("user_request_details_id")]
            # No records for embedding, means embedding is available
            if len(embedding_pending_docs) == 0:
                embeddings = generate_embeddings({1: user_query})
                code_documentation_details = database_utils.get_similarity_searched_documents(embeddings[1], user_request_id)
            
        context_list = []
        for file_name, file_code_content, doc_file_name, documentation_file_content in code_documentation_details:
            # This means its BRD level document so continue
            if file_name is None: continue
            logging.warning(f"file_name: {file_name}")
            # logging.warning(f"file_code_content: {file_code_content}")
            logging.warning(f"doc_file_name: {doc_file_name}")
            # logging.warning(f"documentation_file_content: {documentation_file_content}")
            context_list.append(f"File Name:{file_name}\n\nCode:{file_code_content}\n\nCode Documentation:{documentation_file_content}")
        
        splitted_documentation_content_list = split_documentations(context_list)
        user_query_contents = [] 
        user_query_response = ""             
        for documentation_content in splitted_documentation_content_list:
            user_query_generation_prompt = prompts.get_user_query_prompt(documentation_content, additional_context, user_query)
            user_query_response = generate_oci_gen_ai_response(llm_model, [{"role": "user", "content": user_query_generation_prompt}])
            user_query_contents.append(user_query_response)

        if len(user_query_contents) > 1:
            user_query_content_reponse = "\n".join(user_query_contents)
            document_consolidation_prompt = prompts.get_document_consolidation_prompt(user_query_content_reponse, additional_context)                  
            user_query_response = generate_oci_gen_ai_response(llm_model, [{"role": "user", "content": document_consolidation_prompt}])
        
        # db_user_request = database_utils.db_get_user_request(user_request_id)
        # updated_by = db_user_request["updated_by"]
        # df = pd.DataFrame(columns=['user_request_details_id', 'user_request_id', 'documentation_file_name', 'documentation_file_content', 'success_flag', 'error_details', 'updated_by'])        
        # timestamp = int(time.time())
        # df.loc[0] = [None, user_request_id, f"documentation/UserQuery_{timestamp}.md", user_query_response.replace("## Answer","## Details of your Query"), "Y", None, updated_by]
        # df = generate_docx_from_md(df)
        # database_utils.update_documentation_report(df)
        logging.warning("END: handle_documentation_user_request")
        return user_query_response
    except Exception as e:
        logging.warning(f"Error in handle_documentation_user_request: {str(e)}")
        logging.warning(traceback.format_exc())
        raise


def generate_conversion_report(llm_model, user_request_id, test_mode, additional_context):
    logging.warning("START: generate_conversion_report")
    """
    Converts AngularJS code to ReactJS for a given user request ID.

    Args:
        llm_model (str): The LLM model to use for conversion.
        user_request_id (int): The ID of the user request.
        test_mode (bool): Whether the function is running in test mode.
        additional_context (str): Additional context for the conversion.

    Raises:
        Exception: If there is an error during the conversion process.
    """
    logging.warning(f"START: generate_conversion_report with user_request_id={user_request_id}, test_mode={test_mode}")
    start_time = time.time()
    try:
        db_user_request = database_utils.db_get_user_request(user_request_id)
        if db_user_request["conversion_request_status"] == "In-Progress":
            logging.warning(f"Conversion process is already running for the request ID:{user_request_id}")
            return
        database_utils.update_user_request_status(user_request_id, {"conversion_request_status":"In-Progress","error_message":None})
        package_name = db_user_request["zip_file_name"].split(".")[0]
        updated_by = db_user_request["updated_by"]
        source_language = database_utils.get_language_details(db_user_request["source_lang_id"])
        target_language = database_utils.get_language_details(db_user_request["target_lang_id"])

        logging.warning(f"Package name: {package_name}, Updated by: {updated_by}, Source language: {source_language}, Target language: {target_language}")

        files_to_process, _, file_objects = get_files_from_database(user_request_id, exclude_files=None, include_files=ANGULER_JS_FILE_INCLUSIONS)
        logging.warning(f"Number of files to process: {len(files_to_process)}")

        if len(files_to_process) == 0:
            logging.warning("No files found in the source folder")
            raise Exception("No files found in the source folder")

        if source_language.lower() != "angularjs" or target_language.lower() != "reactjs":
            logging.warning("Invalid source or target language selected.")
            raise Exception("Invalid source or target language selected. Please select AngularJS as source and ReactJS as target language.")

        target_folder_structure = get_files_from_directory(SAMPLE_REACT_PROJECT_PATH, exclude_files=None, include_files=ANGULER_JS_FILE_INCLUSIONS, include_empty_folders=True)
        user_request_details_id, route_file, target_file_to_create, view_controller_info, route_file_content = generate_route_file(llm_model, files_to_process)

        df = pd.DataFrame(columns=['user_request_details_id', 'user_request_id', 'conversion_file_name', 'conversion_file_content', 'success_flag', 'error_details', 'updated_by'])
        data_counter = 0
        errors = []

        if user_request_details_id is not None:
            target_folder_structure.append(route_file)
            del files_to_process[user_request_details_id]
            df.loc[data_counter] = [user_request_details_id, user_request_id, target_file_to_create, route_file_content, "Y", None, updated_by]
            data_counter += 1

            output, errors, target_folder_structure = handle_view_controller_files(llm_model, files_to_process, file_objects, target_folder_structure, view_controller_info, route_file_content, test_mode)
            for item in output:
                view_or_controller_file_name = item[0]
                view_or_controller_target_file_to_create = item[1]
                converted_code = item[2]
                success = item[3]
                error = item[4]
                for obj in file_objects:
                    if obj["file_name"] == view_or_controller_file_name:
                        view_or_controller_id = obj["user_request_details_id"]
                        df.loc[data_counter] = [view_or_controller_id, user_request_id, view_or_controller_target_file_to_create, converted_code, success, error, updated_by]
                        data_counter += 1
                        del files_to_process[view_or_controller_id]
                        break

        logging.warning(f"After processing route file, remaining files to process: {len(files_to_process)}")
        files_to_not_override = ["App.jsx", "main.jsx", "index.html", "package.json", "vite.config.js"]

        for user_request_details_id, (file_name, file_content) in files_to_process.items():
            try:
                logging.warning(f"Processing file: {file_name}")
                if "/tests/" in file_name:
                    logging.warning(f"Skipping test file: {file_name}")
                    continue

                extension = file_name.split(".")[-1]
                f_name = file_name.split("/")[-1]
                files_to_copy = ["png", "jpg", "jpeg", "css", "xml", "csv", "json"]

                if extension in files_to_copy:
                    target_file_to_create = f"react_project/src/assets/{f_name}"
                    logging.warning(f"Skipping image file temporarily: {file_name}")
                    df.loc[data_counter] = [user_request_details_id, user_request_id, target_file_to_create, "SOURCE_FILE_CONTENT", "Y", None, updated_by]
                    data_counter += 1
                else:
                    code_converion_prompt = prompts.get_ang_to_react_code_conversion_prompt(files_to_process, file_name, file_content, additional_context)
                    response = generate_oci_gen_ai_response(llm_model, [{"role": "user", "content": code_converion_prompt}])
                    target_file_name, target_folder_path, converted_code = __get_details_from_repsonse(response)

                    if target_file_name in files_to_not_override:
                        target_file_name = f"{target_file_name}_{data_counter}.jsx"

                    target_file_to_create = SAMPLE_REACT_PROJECT_PATH + "/" + target_folder_path + "/" + target_file_name
                    target_file_to_create = target_file_to_create.replace("\\", "/").replace("//", "/").replace("/\\", "/").replace("react_project/react_project", "react_project")
                    converted_code = converted_code.replace("```javascript", "").replace("```", "").replace("```jsx", "").replace("```", "").replace("jsx", "")

                    df.loc[data_counter] = [user_request_details_id, user_request_id, target_file_to_create, converted_code, "Y", None, updated_by]
                    data_counter += 1
            except Exception as e:
                error = str(e)
                logging.warning(f"Error processing file {file_name}: {error}")
                logging.warning(traceback.format_exc())
                df.loc[data_counter] = [user_request_details_id, user_request_id, None, None, "N", error, updated_by]
                data_counter += 1
                errors.append(file_name)

            if test_mode and data_counter >= TEST_SIZE:
                logging.warning("Test mode enabled, stopping after processing 3 files.")
                break
        
        # Add files from react_project folder into df entries
        react_project_files = get_files_from_directory(SAMPLE_REACT_PROJECT_PATH, exclude_files=None, include_files=None, include_empty_folders=False)
        for file_path in react_project_files:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    file_content = f.read()
                rel_path = file_path.replace("\\", "/")
                df.loc[data_counter] = [None, user_request_id, rel_path, file_content, "Y", None, updated_by]
                data_counter += 1
            except Exception as e:
                error = str(e)
                logging.warning(f"Error reading react_project file {file_path}: {error}")
                logging.warning(traceback.format_exc())
                df.loc[data_counter] = [None, user_request_id, file_path, None, "N", error, updated_by]
                data_counter += 1
        logging.warning(f"Errors encountered: {errors}")
        logging.warning(f"Completed converting code from: {source_language} to {target_language} for package_name: {package_name}")
        database_utils.update_conversion_report(df)
        database_utils.update_user_request_status(user_request_id, {"conversion_request_status":"Done","error_message":None})
    except Exception as e:
        logging.warning(f"Error in generate_conversion_report: {str(e)}")
        logging.warning(traceback.format_exc())
        database_utils.update_user_request_status(user_request_id, {"conversion_request_status":"Done", "error_message":str(e)})
        raise
    finally:
        end_time = time.time()
        total_time = end_time - start_time
        logging.warning(f"END: generate_conversion_report with user_request_id={user_request_id}, test_mode={test_mode}. Total time taken: {total_time:.2f} seconds")